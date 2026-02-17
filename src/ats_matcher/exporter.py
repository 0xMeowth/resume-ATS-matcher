from __future__ import annotations

from typing import Dict

from docx import Document

from ats_matcher.models import ResumeData


class Exporter:
    def apply_changes(
        self, docx_bytes: bytes, resume: ResumeData, changes: Dict[str, str]
    ) -> bytes:
        from io import BytesIO

        document = Document(BytesIO(docx_bytes))
        for bullet_id, new_text in changes.items():
            bullet = resume.bullet_index.get(bullet_id)
            if not bullet:
                continue
            if bullet.paragraph_index >= len(document.paragraphs):
                continue
            paragraph = document.paragraphs[bullet.paragraph_index]
            paragraph.text = new_text
            bullet.text = new_text

        output = self._to_bytes(document)
        return output

    def _to_bytes(self, document: Document) -> bytes:
        from io import BytesIO

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
